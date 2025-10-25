import os
from datetime import datetime
from typing import Dict, Tuple, Optional

from docxtpl import DocxTemplate
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT


def ensure_template(template_path: str) -> None:
    os.makedirs(os.path.dirname(template_path), exist_ok=True)
    if os.path.exists(template_path):
        return
    # Create a default template DOCX with placeholders
    doc = Document()
    doc.add_heading('Metrica Internship Assignment - Document Template', level=1)
    fields = [
        ("Full Name", "{{FullName}}"),
        ("Email Address", "{{Email}}"),
        ("Mobile Number", "{{Mobile}}"),
        ("Company / Institute Name", "{{Company}}"),
        ("Department / Role", "{{Role}}"),
        ("Address", "{{Address}}"),
        ("City", "{{City}}"),
        ("State", "{{State}}"),
        ("Pin Code", "{{PinCode}}"),
        ("Date of Submission", "{{Date}}"),
        ("Remarks / Notes", "{{Remarks}}"),
    ]
    for label, placeholder in fields:
        doc.add_paragraph(f"{label}: {placeholder}")
    doc.save(template_path)


def render_docx(template_path: str, context: Dict[str, str], out_path: str) -> None:
    ensure_template(template_path)
    tpl = DocxTemplate(template_path)
    tpl.render(context)
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    tpl.save(out_path)


def try_convert_to_pdf(docx_path: str, pdf_path: str, context: Dict[str, str]) -> Tuple[bool, Optional[str]]:
    """
    Generate PDF using ReportLab (pure Python, no external dependencies).
    Falls back to returning False if generation fails.
    """
    try:
        # Create PDF directly using the context data
        doc = SimpleDocTemplate(pdf_path, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=20,
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=12,
            spaceAfter=8,
            textColor='#333333',
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=12,
        )
        
        # Title
        story.append(Paragraph("Metrica Internship Assignment - Document Template", title_style))
        story.append(Spacer(1, 0.2 * inch))
        
        
        # Form fields
        fields = [
            ("Full Name", context.get("FullName", "")),
            ("Email Address", context.get("Email", "")),
            ("Mobile Number", context.get("Mobile", "")),
            ("Company / Institute Name", context.get("Company", "")),
            ("Department / Role", context.get("Role", "")),
            ("Address", context.get("Address", "")),
            ("City", context.get("City", "")),
            ("State", context.get("State", "")),
            ("Pin Code", context.get("PinCode", "")),
            ("Date of Submission", context.get("Date", "")),
            ("Remarks / Notes", context.get("Remarks", "")),
        ]
        
        for label, value in fields:
            story.append(Paragraph(f"<b>{label}:</b> {value or 'N/A'}", body_style))
        
        # Build PDF
        doc.build(story)
        return True, None
    except Exception as e:
        return False, str(e)


def build_context(form: Dict[str, Optional[str]]) -> Dict[str, str]:
    # Normalize and default values
    def val(k: str) -> str:
        v = form.get(k)
        return v if v is not None else ""

    # If Date missing, use today
    date_str = form.get("Date") or datetime.now().strftime("%Y-%m-%d")

    return {
        "FullName": val("FullName"),
        "Email": val("Email"),
        "Mobile": val("Mobile"),
        "Company": val("Company"),
        "Role": val("Role"),
        "Address": val("Address"),
        "City": val("City"),
        "State": val("State"),
        "PinCode": val("PinCode"),
        "Date": date_str,
        "Remarks": val("Remarks"),
    }
